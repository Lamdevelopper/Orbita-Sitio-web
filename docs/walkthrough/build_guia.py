from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "docs" / "walkthrough" / "Guia_entrega_articulos_Orbita.docx"

NAVY = "17255B"
BLUE = "506AB8"
LIGHT_BLUE = "E8EEF9"
PALE = "F5F7FB"
INK = "171B23"
MUTED = "5F6878"
RED = "C8463A"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_paragraph(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        borders.append(left)
        p_pr.append(borders)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_para(doc, text="", *, bold_lead=None, italic=False, color=INK, size=11, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, size=size, color=color, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=size, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_numbered(doc, lead, detail):
    p = doc.add_paragraph(style="List Number")
    first = p.add_run(lead + " ")
    set_run_font(first, bold=True)
    rest = p.add_run(detail)
    set_run_font(rest)
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_code(doc, text, label=None):
    if label:
        add_para(doc, label, color=BLUE, size=9.5, after=3)
    p = doc.add_paragraph(style="Code Block")
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Courier New", size=8.7, color=INK)
    return p


def add_callout(doc, label, text, tone="blue"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, LIGHT_BLUE if tone == "blue" else "FCEEEB", BLUE if tone == "blue" else RED)
    lead = p.add_run(label.upper() + "  ")
    set_run_font(lead, size=9.5, color=BLUE if tone == "blue" else RED, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=INK)
    return p


def add_page_break(doc):
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13.5)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(20)

    for list_name in ("List Number", "List Bullet"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in [style.name for style in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    code.font.name = "Courier New"
    code.font.size = Pt(8.7)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(10)
    code.paragraph_format.line_spacing = 1.0

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("ÓRBITA  ·  GUÍA DE COLABORACIÓN")
    set_run_font(hr, size=8.5, color=BLUE, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Órbita · Aerospace AAFI     ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def build():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Guía para entregar artículos a Órbita"
    doc.core_properties.subject = "Plantilla editorial para escritores"
    doc.core_properties.author = "Equipo editorial de Órbita"

    # Page 1: operational customer-pack opening.
    kicker = add_para(doc, "GUÍA DE COLABORACIÓN", color=RED, size=10, after=4)
    kicker.runs[0].bold = True
    title = doc.add_paragraph(style="Title")
    title.add_run("Cómo entregar tu artículo a Órbita")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Una carpeta, una plantilla sencilla y todas tus imágenes listas para publicar.")
    add_callout(doc, "Lo más importante", "No necesitas saber de páginas web. Entrega una carpeta ordenada; el equipo editorial se encarga de revisar, dar formato y publicar.")
    add_heading(doc, "El proceso en cinco pasos", 1)
    add_numbered(doc, "Crea una carpeta.", "Ponle un nombre corto relacionado con tu artículo.")
    add_numbered(doc, "Copia la plantilla.", "Guárdala como articulo.txt o articulo.md dentro de la carpeta.")
    add_numbered(doc, "Escribe sin borrar las etiquetas.", "Completa TÍTULO, AUTOR, CATEGORÍA, BAJADA y, si aplica, EDICIÓN.")
    add_numbered(doc, "Agrega tus imágenes.", "Guarda los archivos originales en la misma carpeta y escribe su nombre exacto en RUTA.")
    add_numbered(doc, "Comprime y envía.", "Convierte la carpeta en un archivo .zip y mándalo al responsable editorial.")
    add_heading(doc, "Qué hará el editor", 2)
    add_bullet(doc, "Revisará el texto y abrirá una vista previa antes de publicar.")
    add_bullet(doc, "Subirá las imágenes y conservará los pies de foto que escribiste.")
    add_bullet(doc, "Te avisará si falta un dato o si una imagen necesita mayor resolución.")
    add_para(doc, "Tiempo para preparar la entrega: normalmente 5 a 10 minutos después de terminar el texto.", italic=True, color=MUTED, size=10)

    # Page 2: folder and exact copyable template.
    add_page_break(doc)
    add_heading(doc, "1. Prepara una sola carpeta", 1)
    add_para(doc, "La carpeta evita que se pierdan imágenes, nombres o leyendas. Todo lo necesario para publicar debe quedar dentro.")
    add_code(doc, "mi-articulo-cansat/\n  articulo.txt\n  foto-del-evento.webp\n  equipo-en-lanzamiento.jpg", "EJEMPLO DE CARPETA")
    add_callout(doc, "Regla de nombres", "Usa letras, números y guiones. Evita nombres como IMG_0042 final FINAL.jpg. Prefiere foto-del-evento.jpg.")
    add_heading(doc, "2. Copia esta plantilla", 1)
    template = """TÍTULO:
AUTOR:
CATEGORÍA:
BAJADA:
EDICIÓN:

---

## Primer subtítulo

Escribe aquí tu primer párrafo.

> Una cita breve que quieras destacar.

## Segundo subtítulo

Continúa aquí el artículo.

[IMAGEN 1]
RUTA: nombre-exacto-de-la-imagen.jpg
PIE DE FOTO: Explica qué aparece, dónde y quién tomó la foto.

## Conclusión

Cierra con la idea que quieres dejar en tus lectores."""
    add_code(doc, template)
    add_para(doc, "Guarda el archivo como articulo.txt o articulo.md. No hace falta cambiar tipografías, colores ni márgenes.", italic=True, color=MUTED, size=10)

    # Page 3: writing syntax and images.
    add_page_break(doc)
    add_heading(doc, "3. Escribe el artículo", 1)
    add_heading(doc, "Datos de la parte superior", 2)
    add_bullet(doc, "TÍTULO: obligatorio. Usa el título que quieres que vea el lector.", "TÍTULO:")
    add_bullet(doc, "AUTOR: obligatorio. Escribe tu nombre completo.", "AUTOR:")
    add_bullet(doc, "CATEGORÍA: obligatoria. Ejemplos: Ingeniería, Entrevista, Aeroespacial, Investigación o Comunidad.", "CATEGORÍA:")
    add_bullet(doc, "BAJADA: una o dos frases que resumen la historia y despiertan interés.", "BAJADA:")
    add_bullet(doc, "EDICIÓN: opcional. Déjala vacía si el editor no te indicó una edición.", "EDICIÓN:")
    add_callout(doc, "No borres esta línea", "El separador --- debe ir solo en una línea entre los datos y el cuerpo del artículo.", tone="red")
    add_heading(doc, "Señales sencillas dentro del texto", 2)
    add_code(doc, "## Un subtítulo\n\n> Una cita destacada")
    add_bullet(doc, "Dos signos ## convierten esa línea en subtítulo.")
    add_bullet(doc, "El signo > al inicio convierte la frase en una cita destacada.")
    add_bullet(doc, "Deja una línea vacía entre párrafos para que el texto sea fácil de revisar.")
    add_heading(doc, "Imágenes y pies de foto", 1)
    add_para(doc, "Coloca el bloque exactamente en el lugar donde quieres que aparezca la imagen. Cada dato ocupa una sola línea.")
    add_code(doc, "[IMAGEN 1]\nRUTA: foto-del-evento.webp\nPIE DE FOTO: El equipo prepara el CanSat antes del lanzamiento. Foto: Ana López.")
    add_bullet(doc, "Numera en orden: [IMAGEN 1], [IMAGEN 2], [IMAGEN 3]...")
    add_bullet(doc, "RUTA debe coincidir exactamente con el archivo dentro de la carpeta.")
    add_bullet(doc, "El pie debe explicar qué ocurre, identificar personas relevantes y dar crédito cuando corresponda.")
    add_bullet(doc, "Envía JPG, PNG o WebP en la mayor resolución disponible. No pegues capturas dentro del documento.")

    # Page 4: complete example.
    add_page_break(doc)
    add_heading(doc, "4. Ejemplo completo", 1)
    add_para(doc, "Este ejemplo puede pasar directamente por la vista previa del editor.")
    example = """TÍTULO: La primera misión de nuestro equipo CanSat
AUTOR: Andrea Pérez López
CATEGORÍA: Ingeniería
BAJADA: Diseñar un satélite del tamaño de una lata nos enseñó a convertir límites reales en decisiones de ingeniería.
EDICIÓN: julio-2026

---

## Una misión pequeña con preguntas grandes

Nuestro equipo comenzó con una pregunta sencilla: ¿cuántos sistemas podíamos integrar dentro del volumen de una lata sin perder confiabilidad?

Durante tres meses diseñamos la alimentación, la telemetría y el sistema de recuperación. Cada prueba dejó datos y una lista clara de cambios.

> El prototipo mejoró cuando dejamos de ocultar los errores y empezamos a documentarlos.

## El día del lanzamiento

La mañana del lanzamiento revisamos conexiones, batería y recepción de datos. El descenso fue estable y recuperamos el CanSat sin daños.

[IMAGEN 1]
RUTA: equipo-en-lanzamiento.jpg
PIE DE FOTO: El equipo verifica la telemetría minutos antes del lanzamiento. Foto: Andrea Pérez.

## Lo que sigue

La siguiente versión incorporará un sensor ambiental y una carcasa más ligera. También publicaremos el registro de pruebas para que otros equipos puedan aprender del proceso."""
    add_code(doc, example)

    # Page 5: final checks and handoff.
    add_page_break(doc)
    add_heading(doc, "5. Revisa antes de enviar", 1)
    checklist = [
        "El archivo incluye TÍTULO, AUTOR y CATEGORÍA.",
        "La línea --- aparece entre los datos y el artículo.",
        "Los subtítulos comienzan con ##.",
        "Cada imagen está dentro de la carpeta y tiene un bloque numerado.",
        "Cada RUTA coincide letra por letra con el nombre del archivo.",
        "Cada imagen tiene PIE DE FOTO y crédito cuando corresponde.",
        "Revisaste ortografía, nombres propios, cifras y enlaces.",
        "La carpeta abre correctamente después de comprimirla como .zip.",
    ]
    for item in checklist:
        add_bullet(doc, item)
    add_heading(doc, "Cómo enviar la carpeta", 1)
    add_numbered(doc, "Comprímela.", "En Windows: clic derecho sobre la carpeta > Comprimir en archivo ZIP. En macOS: clic derecho > Comprimir.")
    add_numbered(doc, "Nombra el ZIP.", "Usa tu apellido y un título corto, por ejemplo: perez-mision-cansat.zip.")
    add_numbered(doc, "Envíalo.", "Adjunta el ZIP por el medio acordado con el responsable editorial y escribe tu nombre y título en el mensaje.")
    add_callout(doc, "Si el archivo pesa demasiado", "Comparte una carpeta de Drive con permiso de lectura y no muevas ni renombres archivos después de enviar el enlace.")
    add_heading(doc, "Lo que no necesitas hacer", 2)
    add_bullet(doc, "No necesitas diseñar la página ni acomodar tamaños para web.")
    add_bullet(doc, "No necesitas subir nada al sitio ni tener acceso al panel administrativo.")
    add_bullet(doc, "No necesitas convertir las imágenes: envía los originales y el editor resolverá la publicación.")
    add_para(doc, "¿Tienes dudas? Pregunta antes de cambiar la plantilla. Una pregunta breve suele ahorrar una ronda de correcciones.", bold_lead="¿Tienes dudas?", color=NAVY, after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
